from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "artifacts" / "ThisCafeteria_AWS_Wallet_Status_Runbook.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D3DF"
CODE_FILL = "F5F7FA"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Pt(width / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if keep_next is None:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.font.bold = True
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(80, 90, 105)
    subtitle.paragraph_format.space_after = Pt(12)

    code = styles.add_style("Code Block", 1)
    code.font.name = "Courier New"
    code.font.size = Pt(8.5)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.05)
    code.paragraph_format.line_spacing = 1.0

    callout = styles.add_style("Callout", 1)
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.left_indent = Inches(0.15)
    callout.paragraph_format.right_indent = Inches(0.15)
    callout.paragraph_format.line_spacing = 1.18


def add_code(doc: Document, text: str):
    p = doc.add_paragraph(style="Code Block")
    run = p.add_run(text.rstrip())
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    p_format = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_FILL)
    p_format.append(shd)
    return p


def add_note(doc: Document, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.style = doc.styles["Callout"]
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run(text)
    return table


def add_kv_table(doc: Document, rows, widths=(2400, 6960)):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, list(widths))
    hdr = table.rows[0]
    hdr.cells[0].text = "Item"
    hdr.cells[1].text = "Value"
    set_repeat_table_header(hdr)
    for cell in hdr.cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    return table


def add_three_table(doc: Document, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(headers):
        hdr.cells[i].text = head
        set_cell_shading(hdr.cells[i], LIGHT_BLUE)
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    return table


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("ThisCafeteria AWS Deployment Documentation")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(80, 90, 105)

    title = doc.add_paragraph(style="Title")
    title.add_run("ThisCafeteria AWS Wallet Status and EC2 Deployment Runbook")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Session documentation covering the backend status pipeline, AWS resources, "
        "RDS/PostgreSQL, SQS publishing, and EC2 deployment path."
    )

    add_kv_table(
        doc,
        [
            ("Project root", "/Users/alexis/TCDE/ThisCafeteria"),
            ("Document date", "May 26, 2026"),
            ("AWS region", "us-east-1"),
            ("Current deployment target", "EC2 Ubuntu instance with Nginx reverse proxy"),
            ("Application type", "ASP.NET Core web app with Blazor Server/Razor Components and API controllers"),
            ("Status pipeline", "POST /api/wallet-status -> RDS PostgreSQL -> SQS wallet-status"),
        ],
    )

    add_note(
        doc,
        "Secrets policy",
        "No database password, AWS access key, private key, or wallet private key is recorded in this document. "
        "Use user-secrets locally and /etc/thiscafeteria/thiscafeteria.env on EC2.",
    )
    doc.add_page_break()


def build_doc():
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This session built and validated a durable wallet-login/status pipeline for the ThisCafeteria project. "
        "The current production-style architecture is a single ASP.NET Core web application serving the frontend "
        "and API from EC2, with PostgreSQL on RDS as the system of record and SQS as the event stream."
    )
    add_note(
        doc,
        "Current working result",
        "The local backend successfully stored wallet status rows in RDS and published to SQS, returning "
        "publishedToAws=true with an AWS message id. The app package was then copied to EC2, installed as a "
        "systemd service, and verified locally through Kestrel and Nginx health checks.",
    )

    doc.add_heading("2. Application Classification", level=1)
    add_kv_table(
        doc,
        [
            ("Framework", "ASP.NET Core web app"),
            ("UI model", "Blazor Server / Razor Components"),
            ("API surface", "ASP.NET Core MVC API controllers"),
            ("Database access", "Entity Framework Core with Npgsql provider"),
            ("AWS messaging", "AWS SDK for .NET, SQS SendMessage"),
            ("Not this app", "Not a static-only frontend; not Next.js for the deployed path"),
        ],
    )
    doc.add_paragraph(
        "Because the UI is Blazor Server, the frontend cannot be hosted directly on S3 as-is. It needs a running "
        ".NET server for circuits, API routes, wallet auth, database writes, and SQS publishing. Therefore the "
        "selected AWS Free Tier-oriented path is EC2 plus Nginx."
    )

    doc.add_heading("3. Architecture", level=1)
    add_code(
        doc,
        """
Browser
  -> http://EC2_PUBLIC_IP
  -> Nginx on EC2 port 80
  -> Kestrel on 127.0.0.1:5000
  -> ThisCafeteria.Web
       -> EF Core / Npgsql -> RDS PostgreSQL
       -> AWS SDK SQS     -> wallet-status queue
""",
    )
    doc.add_paragraph(
        "The frontend never connects directly to RDS or SQS. Browser code calls the backend API only. "
        "The backend owns database credentials, AWS credentials through IAM, and all durable status publishing."
    )

    doc.add_heading("4. AWS Infrastructure Inventory", level=1)
    add_three_table(
        doc,
        ["Resource", "Value", "Purpose"],
        [
            ("Region", "us-east-1", "All resources used during this session."),
            ("Account ID", "419197236352", "AWS account used in ARNs and queue URLs."),
            ("RDS endpoint", "thiscafeteria.ce3wcicu69fo.us-east-1.rds.amazonaws.com", "PostgreSQL database host."),
            ("RDS port", "5432", "PostgreSQL TCP port."),
            ("RDS database", "thiscafeteria", "Application database."),
            ("RDS username", "appuser", "Application database user."),
            ("SQS queue", "wallet-status", "Queue receiving wallet status events."),
            ("SQS queue URL", "https://sqs.us-east-1.amazonaws.com/419197236352/wallet-status", "Runtime QueueUrl used by backend."),
            ("SQS queue ARN", "arn:aws:sqs:us-east-1:419197236352:wallet-status", "IAM policy resource."),
            ("IAM policy", "wallet-status-send-message-dev", "Allows sqs:SendMessage to wallet-status."),
            ("EC2 instance id", "i-06361d2f4fc6d44e3", "Ubuntu instance created for web deployment."),
            ("EC2 public IP", "3.227.24.88", "Current public HTTP endpoint. Subject to change unless Elastic IP is attached."),
            ("EC2 key pair", "thiscafeteria-dev-key", "SSH access key on local Mac."),
            ("EC2 service path", "/opt/thiscafeteria/ThisCafeteria.Web", "Published self-contained app binary."),
            ("systemd service", "thiscafeteria.service", "Keeps app running on EC2."),
            ("Nginx site", "/etc/nginx/sites-available/thiscafeteria", "Reverse proxy from port 80 to 127.0.0.1:5000."),
        ],
        [1800, 3960, 3600],
    )

    doc.add_heading("5. Free Tier Cost Guardrails", level=1)
    doc.add_paragraph(
        "The deployment aims at AWS Free Tier-style usage, but free eligibility depends on account age, region, "
        "instance classes, storage size, and monthly hours. Billing should be monitored with a budget alert."
    )
    add_three_table(
        doc,
        ["Service", "Free-tier-oriented choice", "Notes"],
        [
            ("EC2", "t3.micro Ubuntu, 8 GB gp3 EBS", "Free Tier commonly covers one micro instance for 750 hours/month if eligible."),
            ("EBS", "8 GB root volume", "Kept below the common 30 GB Free Tier EBS allowance."),
            ("RDS", "Single-AZ db.t3.micro or db.t4g.micro PostgreSQL", "AWS RDS Free Tier includes 750 hours/month for eligible micro PostgreSQL instances."),
            ("SQS", "Standard queue", "SQS provides 1 million requests/month free; this project uses tiny dev traffic."),
            ("App Runner", "Not selected", "Easier managed deployment, but not treated as Free Tier for this runbook."),
        ],
        [1900, 3300, 4160],
    )
    p = doc.add_paragraph("References: ")
    add_hyperlink(p, "EC2 Free Tier", "https://docs.aws.amazon.com/AWSEC2/latest/UserGuide/ec2-free-tier-usage.html")
    p.add_run("; ")
    add_hyperlink(p, "RDS Free Tier", "https://aws.amazon.com/rds/free/")
    p.add_run("; ")
    add_hyperlink(p, "SQS pricing", "https://aws.amazon.com/sqs/pricing/")
    p.add_run("; ")
    add_hyperlink(p, "EBS pricing", "https://aws.amazon.com/ebs/pricing/")

    doc.add_heading("6. Code Changes Implemented", level=1)
    add_three_table(
        doc,
        ["Area", "Files", "What changed"],
        [
            (
                "Wallet status API",
                "src/ThisCafeteria.Web/Controllers/WalletStatusController.cs; src/ThisCafeteria.Web/Models/WalletAuthModels.cs",
                "Added POST /api/wallet-status and GET /api/wallet-status/{walletAddress}. The POST validates walletAddress and status, stores a row, publishes to SQS, and returns the DB id plus AWS message id.",
            ),
            (
                "Domain and EF model",
                "src/ThisCafeteria.Domain/Entities/WalletStatusEvent.cs; Infrastructure/Persistence/Configurations/WalletStatusEventConfiguration.cs",
                "Added wallet status event entity mapped to wallet_status_events with payload_json jsonb, created_at, published_to_aws_at, and aws_message_id.",
            ),
            (
                "Repository",
                "Application/Repositories/IWalletStatusEventRepository.cs; Infrastructure/Persistence/Repositories/WalletStatusEventRepository.cs",
                "Added methods to insert events, read latest event for wallet, read recent events, and mark event published with AWS message id.",
            ),
            (
                "Migration",
                "Infrastructure/Persistence/Migrations/20260526000000_AddWalletStatusEvents*.cs; AppDbContextModelSnapshot.cs",
                "Added EF migration and snapshot metadata so dotnet ef database update creates wallet_status_events correctly.",
            ),
            (
                "AWS SQS",
                "Infrastructure/Configuration/AwsMessagingOptions.cs; Infrastructure/Services/SqsMessagePublisher.cs; Infrastructure/DependencyInjection.cs",
                "Configured SQS client from AWS_REGION, SQS_QUEUE_URL, optional AWS_PROFILE, or IAM role. PublishAsync returns AWS MessageId.",
            ),
            (
                "Database configuration",
                "Infrastructure/Configuration/DatabaseConnectionStringFactory.cs",
                "Builds Npgsql connection string from DB_HOST, DB_PORT, DB_NAME, DB_USERNAME, DB_PASSWORD, with SSL Mode=Require and Trust Server Certificate=true.",
            ),
            (
                "Wallet login integration",
                "Web/Controllers/WalletAuthController.cs; Web/wwwroot/js/walletAuth.js; Components/Layout/NavMenu.razor",
                "Wallet auth challenge/verify/logout now record and publish wallet status events in the same RDS + SQS flow.",
            ),
            (
                "Docs",
                "docs/aws-wallet-status.md",
                "Created a focused implementation guide for RDS, SQS, user-secrets, migrations, curl tests, and frontend API usage.",
            ),
        ],
        [1500, 3000, 4860],
    )

    doc.add_heading("7. Database Design", level=1)
    add_code(
        doc,
        """
CREATE TABLE wallet_status_events (
    id uuid PRIMARY KEY,
    wallet_address text NOT NULL,
    status text NOT NULL,
    event_type text NULL,
    payload_json jsonb NULL,
    created_at timestamptz NOT NULL DEFAULT now(),
    published_to_aws_at timestamptz NULL,
    aws_message_id text NULL
);

CREATE INDEX ix_wallet_status_events_wallet_created_at
    ON wallet_status_events (wallet_address, created_at);

CREATE INDEX "IX_wallet_status_events_status"
    ON wallet_status_events (status);

CREATE INDEX "IX_wallet_status_events_aws_message_id"
    ON wallet_status_events (aws_message_id);
""",
    )
    doc.add_paragraph(
        "This table is append-only by design. The latest status is obtained by querying the newest created_at row "
        "for a wallet. That preserves audit history while keeping the current-status query simple."
    )

    doc.add_heading("8. Local Development Configuration", level=1)
    add_code(
        doc,
        """
cd /Users/alexis/TCDE/ThisCafeteria/src/ThisCafeteria.Web

dotnet user-secrets init
dotnet user-secrets set "DB_HOST" "thiscafeteria.ce3wcicu69fo.us-east-1.rds.amazonaws.com"
dotnet user-secrets set "DB_PORT" "5432"
dotnet user-secrets set "DB_NAME" "thiscafeteria"
dotnet user-secrets set "DB_USERNAME" "appuser"
dotnet user-secrets set "DB_PASSWORD" "<local DB password>"
dotnet user-secrets set "AWS_REGION" "us-east-1"
dotnet user-secrets set "SQS_QUEUE_URL" "https://sqs.us-east-1.amazonaws.com/419197236352/wallet-status"
dotnet user-secrets set "AWS_PROFILE" "<local AWS CLI profile if needed>"
""",
    )
    add_note(
        doc,
        "AWS credential behavior",
        "Local development can use AWS_PROFILE and AWS CLI/SSO credentials. EC2 should not use AWS_PROFILE; it should use the attached IAM instance role.",
    )

    doc.add_heading("9. EF Migration Runbook", level=1)
    add_code(
        doc,
        """
cd /Users/alexis/TCDE/ThisCafeteria

dotnet ef migrations list \\
  --project src/ThisCafeteria.Infrastructure \\
  --startup-project src/ThisCafeteria.Web

dotnet ef database update \\
  --project src/ThisCafeteria.Infrastructure \\
  --startup-project src/ThisCafeteria.Web
""",
    )
    doc.add_paragraph(
        "Expected migration name: 20260526000000_AddWalletStatusEvents. If EF says the database is up to date "
        "but the table is missing, inspect migration metadata and the __EFMigrationsHistory table before using manual SQL."
    )

    doc.add_heading("10. Local API Tests", level=1)
    add_code(
        doc,
        """
dotnet run --project src/ThisCafeteria.Web --urls "http://localhost:5295"

curl -i -X POST http://localhost:5295/api/wallet-status \\
  -H "Content-Type: application/json" \\
  -d '{"walletAddress":"0x0000000000000000000000000000000000000000","status":"Connected","eventType":"wallet-login.connected","payload":{"source":"curl"}}'

curl -i http://localhost:5295/api/wallet-status/0x0000000000000000000000000000000000000000

aws sqs receive-message \\
  --region us-east-1 \\
  --queue-url https://sqs.us-east-1.amazonaws.com/419197236352/wallet-status \\
  --max-number-of-messages 1 \\
  --wait-time-seconds 5
""",
    )
    doc.add_paragraph(
        "Successful POST response includes publishedToAws=true and an awsMessageId. A 404 on GET means no row exists "
        "yet for that wallet; POST first."
    )

    doc.add_heading("11. Key Troubleshooting Events From Session", level=1)
    add_three_table(
        doc,
        ["Symptom", "Root cause", "Fix used"],
        [
            (
                "Failed to bind to localhost port 5286 or 5291",
                "A previous dotnet process was still running on that port.",
                "Started the app on another port with --urls, for example http://localhost:5292.",
            ),
            (
                "relation wallet_status_events does not exist",
                "The new status table had not been created in RDS.",
                "Fixed EF migration metadata, ran dotnet ef database update, and verified table-backed POST/GET.",
            ),
            (
                "publishedToAws=false with EC2 metadata credential error",
                ".NET AWS SDK could not find local AWS credentials and fell through to EC2 metadata.",
                "Added AWS_PROFILE support locally and set AWS_PROFILE in user-secrets.",
            ),
            (
                "AWSSDK.Signin could not be found",
                "The local AWS profile used newer AWS login/SSO credential flow requiring the Signin runtime package.",
                "Added AWSSDK.Signin package to ThisCafeteria.Infrastructure.csproj and restored.",
            ),
            (
                "scp could not resolve hostname",
                "EC2_PUBLIC_IP was empty in the current shell.",
                "Re-exported EC2_PUBLIC_IP from aws ec2 describe-instances.",
            ),
            (
                "scp local tarball missing",
                "Command was run from home directory instead of project directory.",
                "Changed to /Users/alexis/TCDE/ThisCafeteria and reran scp.",
            ),
        ],
        [2700, 3300, 3360],
    )

    doc.add_heading("12. AWS CLI EC2 Deployment Commands", level=1)
    doc.add_heading("12.1 Variables", level=2)
    add_code(
        doc,
        """
export AWS_REGION=us-east-1
export ACCOUNT_ID=$(aws sts get-caller-identity --query Account --output text)
export APP=thiscafeteria-web
""",
    )
    doc.add_heading("12.2 EC2 Key Pair", level=2)
    add_code(
        doc,
        """
aws ec2 create-key-pair \\
  --region "$AWS_REGION" \\
  --key-name thiscafeteria-dev-key \\
  --query KeyMaterial \\
  --output text > ~/.ssh/thiscafeteria-dev-key.pem

chmod 400 ~/.ssh/thiscafeteria-dev-key.pem
""",
    )
    doc.add_heading("12.3 Security Group", level=2)
    add_code(
        doc,
        """
export VPC_ID=$(aws ec2 describe-vpcs \\
  --region "$AWS_REGION" \\
  --filters Name=is-default,Values=true \\
  --query 'Vpcs[0].VpcId' \\
  --output text)

export MY_IP=$(curl -s https://checkip.amazonaws.com)/32

export EC2_SG_ID=$(aws ec2 create-security-group \\
  --region "$AWS_REGION" \\
  --group-name thiscafeteria-ec2-dev-sg \\
  --description "ThisCafeteria EC2 dev web access" \\
  --vpc-id "$VPC_ID" \\
  --query GroupId \\
  --output text)

aws ec2 authorize-security-group-ingress \\
  --region "$AWS_REGION" \\
  --group-id "$EC2_SG_ID" \\
  --protocol tcp \\
  --port 22 \\
  --cidr "$MY_IP"

aws ec2 authorize-security-group-ingress \\
  --region "$AWS_REGION" \\
  --group-id "$EC2_SG_ID" \\
  --protocol tcp \\
  --port 80 \\
  --cidr 0.0.0.0/0
""",
    )
    doc.add_heading("12.4 IAM Role for EC2", level=2)
    add_code(
        doc,
        """
cat > ec2-trust-policy.json <<'EOF'
{
  "Version": "2012-10-17",
  "Statement": [
    {
      "Effect": "Allow",
      "Principal": { "Service": "ec2.amazonaws.com" },
      "Action": "sts:AssumeRole"
    }
  ]
}
EOF

aws iam create-role \\
  --role-name thiscafeteria-ec2-role \\
  --assume-role-policy-document file://ec2-trust-policy.json

aws iam attach-role-policy \\
  --role-name thiscafeteria-ec2-role \\
  --policy-arn arn:aws:iam::419197236352:policy/wallet-status-send-message-dev

aws iam create-instance-profile \\
  --instance-profile-name thiscafeteria-ec2-profile

aws iam add-role-to-instance-profile \\
  --instance-profile-name thiscafeteria-ec2-profile \\
  --role-name thiscafeteria-ec2-role
""",
    )
    doc.add_heading("12.5 Launch Free-Tier-Sized EC2", level=2)
    add_code(
        doc,
        """
export SUBNET_ID=$(aws ec2 describe-subnets \\
  --region "$AWS_REGION" \\
  --filters Name=vpc-id,Values="$VPC_ID" Name=default-for-az,Values=true \\
  --query 'Subnets[0].SubnetId' \\
  --output text)

export AMI_ID=$(aws ssm get-parameter \\
  --region "$AWS_REGION" \\
  --name /aws/service/canonical/ubuntu/server/24.04/stable/current/amd64/hvm/ebs-gp3/ami-id \\
  --query Parameter.Value \\
  --output text)

export INSTANCE_ID=$(aws ec2 run-instances \\
  --region "$AWS_REGION" \\
  --image-id "$AMI_ID" \\
  --instance-type t3.micro \\
  --key-name thiscafeteria-dev-key \\
  --security-group-ids "$EC2_SG_ID" \\
  --subnet-id "$SUBNET_ID" \\
  --iam-instance-profile Name=thiscafeteria-ec2-profile \\
  --block-device-mappings '[{"DeviceName":"/dev/sda1","Ebs":{"VolumeSize":8,"VolumeType":"gp3","DeleteOnTermination":true}}]' \\
  --tag-specifications "ResourceType=instance,Tags=[{Key=Name,Value=thiscafeteria-web-dev}]" \\
  --query 'Instances[0].InstanceId' \\
  --output text)

aws ec2 wait instance-running --region "$AWS_REGION" --instance-ids "$INSTANCE_ID"

export EC2_PUBLIC_IP=$(aws ec2 describe-instances \\
  --region "$AWS_REGION" \\
  --instance-ids "$INSTANCE_ID" \\
  --query 'Reservations[0].Instances[0].PublicIpAddress' \\
  --output text)

echo "$EC2_PUBLIC_IP"
""",
    )
    doc.add_heading("12.6 Allow EC2 to Reach RDS", level=2)
    add_code(
        doc,
        """
export RDS_SG_ID=$(aws rds describe-db-instances \\
  --region "$AWS_REGION" \\
  --db-instance-identifier thiscafeteria \\
  --query 'DBInstances[0].VpcSecurityGroups[0].VpcSecurityGroupId' \\
  --output text)

aws ec2 authorize-security-group-ingress \\
  --region "$AWS_REGION" \\
  --group-id "$RDS_SG_ID" \\
  --protocol tcp \\
  --port 5432 \\
  --source-group "$EC2_SG_ID"
""",
    )

    doc.add_heading("13. Build and Copy Application to EC2", level=1)
    add_code(
        doc,
        """
cd /Users/alexis/TCDE/ThisCafeteria

dotnet publish src/ThisCafeteria.Web/ThisCafeteria.Web.csproj \\
  -c Release \\
  -r linux-x64 \\
  --self-contained true \\
  -o publish-ec2

tar -czf thiscafeteria-web.tgz -C publish-ec2 .

scp -i ~/.ssh/thiscafeteria-dev-key.pem \\
  thiscafeteria-web.tgz \\
  ubuntu@"$EC2_PUBLIC_IP":/tmp/
""",
    )

    doc.add_heading("14. EC2 Installation Commands", level=1)
    doc.add_paragraph("SSH into EC2:")
    add_code(doc, 'ssh -i ~/.ssh/thiscafeteria-dev-key.pem ubuntu@"$EC2_PUBLIC_IP"')
    doc.add_paragraph("On the EC2 instance:")
    add_code(
        doc,
        """
sudo apt-get update
sudo apt-get install -y nginx

sudo mkdir -p /opt/thiscafeteria /etc/thiscafeteria
sudo tar -xzf /tmp/thiscafeteria-web.tgz -C /opt/thiscafeteria
sudo chmod +x /opt/thiscafeteria/ThisCafeteria.Web
""",
    )
    doc.add_heading("14.1 EC2 Environment File", level=2)
    add_code(
        doc,
        """
sudo nano /etc/thiscafeteria/thiscafeteria.env
""",
    )
    add_code(
        doc,
        """
ASPNETCORE_ENVIRONMENT=Production
ASPNETCORE_URLS=http://127.0.0.1:5000
AWS_REGION=us-east-1
SQS_QUEUE_URL=https://sqs.us-east-1.amazonaws.com/419197236352/wallet-status
DB_HOST=thiscafeteria.ce3wcicu69fo.us-east-1.rds.amazonaws.com
DB_PORT=5432
DB_NAME=thiscafeteria
DB_USERNAME=appuser
DB_PASSWORD=<set on EC2 only>
""",
    )
    doc.add_heading("14.2 systemd Service", level=2)
    add_code(
        doc,
        """
sudo tee /etc/systemd/system/thiscafeteria.service > /dev/null <<'EOF'
[Unit]
Description=ThisCafeteria Web App
After=network.target

[Service]
WorkingDirectory=/opt/thiscafeteria
ExecStart=/opt/thiscafeteria/ThisCafeteria.Web
Restart=always
RestartSec=10
EnvironmentFile=/etc/thiscafeteria/thiscafeteria.env
User=www-data

[Install]
WantedBy=multi-user.target
EOF

sudo systemctl daemon-reload
sudo systemctl enable thiscafeteria
sudo systemctl start thiscafeteria
sudo systemctl status thiscafeteria --no-pager
""",
    )
    doc.add_paragraph("Observed result during session: service became active and ran /opt/thiscafeteria/ThisCafeteria.Web.")

    doc.add_heading("15. Nginx Reverse Proxy", level=1)
    add_code(
        doc,
        """
sudo tee /etc/nginx/sites-available/thiscafeteria > /dev/null <<'EOF'
server {
    listen 80;
    server_name _;

    location / {
        proxy_pass http://127.0.0.1:5000;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection keep-alive;
        proxy_set_header Host $host;
        proxy_cache_bypass $http_upgrade;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }
}
EOF

sudo ln -sf /etc/nginx/sites-available/thiscafeteria /etc/nginx/sites-enabled/thiscafeteria
sudo rm -f /etc/nginx/sites-enabled/default
sudo nginx -t
sudo systemctl restart nginx
""",
    )
    doc.add_heading("15.1 Health Checks Observed", level=2)
    add_code(
        doc,
        """
curl -i http://127.0.0.1:5000/health
curl -i http://127.0.0.1/health

Observed response for both:
HTTP/1.1 200 OK
Healthy
""",
    )

    doc.add_heading("16. Public Verification Commands", level=1)
    add_code(
        doc,
        """
curl -i http://3.227.24.88/health

curl -i -X POST http://3.227.24.88/api/wallet-status \\
  -H "Content-Type: application/json" \\
  -d '{"walletAddress":"0x0000000000000000000000000000000000000000","status":"Connected","eventType":"ec2.public.test","payload":{"source":"ec2-public"}}'

curl -i http://3.227.24.88/api/wallet-status/0x0000000000000000000000000000000000000000
""",
    )
    doc.add_paragraph(
        "If the public health check fails but the EC2-local health checks pass, inspect the EC2 security group inbound rule for TCP 80."
    )

    doc.add_heading("17. Operating the EC2 App", level=1)
    add_three_table(
        doc,
        ["Task", "Command", "Expected use"],
        [
            ("View service status", "sudo systemctl status thiscafeteria --no-pager -l", "Check whether the .NET app is running."),
            ("View app logs", "sudo journalctl -u thiscafeteria -n 100 --no-pager -l", "Debug startup, DB, SQS, and runtime errors."),
            ("Restart app", "sudo systemctl restart thiscafeteria", "Use after changing env file or replacing app files."),
            ("View Nginx status", "sudo systemctl status nginx --no-pager -l", "Check reverse proxy."),
            ("Test Nginx config", "sudo nginx -t", "Run before restarting Nginx after config edits."),
            ("Restart Nginx", "sudo systemctl restart nginx", "Apply reverse proxy changes."),
        ],
        [1800, 3900, 3660],
    )

    doc.add_heading("18. Update Deployment Runbook", level=1)
    add_numbers(
        doc,
        [
            "Make code changes locally and verify them.",
            "Publish a new self-contained linux-x64 build.",
            "Create a new thiscafeteria-web.tgz archive.",
            "Copy the archive to EC2 with scp.",
            "SSH to EC2, stop the service, replace /opt/thiscafeteria contents, and restart.",
            "Run local EC2 health checks and public health checks.",
        ],
    )
    add_code(
        doc,
        """
# Local Mac
cd /Users/alexis/TCDE/ThisCafeteria
dotnet publish src/ThisCafeteria.Web/ThisCafeteria.Web.csproj -c Release -r linux-x64 --self-contained true -o publish-ec2
tar -czf thiscafeteria-web.tgz -C publish-ec2 .
scp -i ~/.ssh/thiscafeteria-dev-key.pem thiscafeteria-web.tgz ubuntu@"$EC2_PUBLIC_IP":/tmp/

# EC2
sudo systemctl stop thiscafeteria
sudo rm -rf /opt/thiscafeteria/*
sudo tar -xzf /tmp/thiscafeteria-web.tgz -C /opt/thiscafeteria
sudo chmod +x /opt/thiscafeteria/ThisCafeteria.Web
sudo systemctl start thiscafeteria
sudo systemctl status thiscafeteria --no-pager -l
""",
    )

    doc.add_heading("19. Stop or Remove Resources to Control Cost", level=1)
    add_code(
        doc,
        """
# Stop EC2 instance when not testing.
aws ec2 stop-instances --region us-east-1 --instance-ids i-06361d2f4fc6d44e3

# Start it again later.
aws ec2 start-instances --region us-east-1 --instance-ids i-06361d2f4fc6d44e3

# Terminate completely when done.
aws ec2 terminate-instances --region us-east-1 --instance-ids i-06361d2f4fc6d44e3
""",
    )
    add_note(
        doc,
        "Cost note",
        "Stopping EC2 stops instance-hour charges but EBS root volume storage may continue billing. Terminating with DeleteOnTermination=true removes the root volume.",
    )

    doc.add_heading("20. Security and Production Hardening", level=1)
    add_bullets(
        doc,
        [
            "Move the public site to HTTPS with a real domain and TLS certificate.",
            "Avoid RDS public access for production; allow inbound PostgreSQL only from the EC2 security group.",
            "Store DB_PASSWORD in AWS Secrets Manager or SSM Parameter Store instead of a plain env file when moving beyond dev.",
            "Protect POST /api/wallet-status. In dev it accepts direct status posts; production should require a verified session or only write statuses inside wallet-auth flows.",
            "Use a real non-zero wallet address for functional wallet tests. The zero address is useful for curl smoke tests only.",
            "Add CloudWatch alarms for EC2 CPU, status check failures, and billing.",
            "Add an SQS dead-letter queue before introducing consumers.",
            "Back up RDS and document restore procedure.",
        ],
    )

    doc.add_heading("21. Next Engineering Steps", level=1)
    add_three_table(
        doc,
        ["Priority", "Step", "Why it matters"],
        [
            ("P0", "Verify public EC2 URL from Mac: http://3.227.24.88/health", "Confirms security group and Nginx are reachable externally."),
            ("P0", "POST /api/wallet-status against public EC2 URL", "Confirms public request -> RDS -> SQS using EC2 IAM role."),
            ("P1", "Test actual wallet login in browser", "Confirms frontend wallet modal drives the backend status pipeline."),
            ("P1", "Change chain config to Base Sepolia", "Moves wallet network from the current BSC testnet config to the target chain."),
            ("P1", "Lock down status writes", "Prevents arbitrary clients from writing fake statuses."),
            ("P2", "Add HTTPS and domain", "Makes the public URL usable and secure for normal browser wallet flows."),
            ("P2", "Add SQS consumer/worker", "Turns published events into downstream processing instead of just queue storage."),
        ],
        [900, 3900, 4560],
    )

    doc.add_heading("22. Appendix: Commands That Proved the Pipeline", level=1)
    add_code(
        doc,
        """
# Local successful status post after AWSSDK.Signin/profile fix.
curl -i -X POST http://localhost:5295/api/wallet-status \\
  -H "Content-Type: application/json" \\
  -d '{"walletAddress":"0x0000000000000000000000000000000000000000","status":"Connected","eventType":"wallet-login.connected","payload":{"source":"curl"}}'

# Observed response:
HTTP/1.1 200 OK
{
  "id": "d48f09c2-4e8a-4670-944d-988aa44af3c2",
  "walletAddress": "0x0000000000000000000000000000000000000000",
  "status": "Connected",
  "eventType": "wallet-login.connected",
  "publishedToAws": true,
  "awsMessageId": "c7d87806-2c95-4ed8-b5dd-dafbf8b551e7"
}

# EC2-local health checks after systemd and Nginx setup:
curl -i http://127.0.0.1:5000/health
curl -i http://127.0.0.1/health

# Observed response:
HTTP/1.1 200 OK
Healthy
""",
    )

    # Footer
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("ThisCafeteria AWS Runbook")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 110, 120)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
